import os
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.schema import Block, Exercise, TrainingSessionPlan
from app.storage import _slug

OUTPUTS_DIR = Path(os.environ.get("OUTPUTS_DIR", str(Path(__file__).parent.parent / "outputs")))


def _output_path(plan: TrainingSessionPlan) -> Path:
    meta = plan.meta
    client_slug = _slug(meta.client_name or "unknown")
    session_num = meta.session_number or 0
    filename = f"{meta.session_date or 'undated'}_session_{session_num}.docx"
    return OUTPUTS_DIR / client_slug / filename


def _machine_parts(ms) -> list[str]:
    parts = []
    if ms.machine_name:
        parts.append(ms.machine_name)
    if ms.seat:
        parts.append(f"Seat {ms.seat}" if not ms.seat.lower().startswith("seat") else ms.seat)
    if ms.lever:
        parts.append(f"Lever {ms.lever}" if not ms.lever.lower().startswith("lever") else ms.lever)
    if ms.pad:
        parts.append(f"Pad {ms.pad}" if not ms.pad.lower().startswith("pad") else ms.pad)
    return parts


def _bold_para(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)


def _add_exercise(doc: Document, ex: Exercise) -> None:
    doc.add_heading(ex.name, level=3)

    details = []
    if ex.sets is not None:
        details.append(f"Sets: {ex.sets}")
    if ex.reps:
        details.append(f"Reps: {ex.reps}")
    if ex.tempo:
        details.append(f"Tempo: {ex.tempo}")
    if ex.rest_seconds is not None:
        details.append(f"Rest: {ex.rest_seconds}s")
    if ex.intensity:
        details.append(f"Intensity: {ex.intensity}")
    if details:
        doc.add_paragraph(" | ".join(details))

    if ex.machine_settings:
        ms = ex.machine_settings
        parts = _machine_parts(ms)
        if parts:
            _bold_para(doc, "Machine", " | ".join(parts))
        if ms.notes:
            _bold_para(doc, "Setup notes", ms.notes)

    if ex.loading:
        ld = ex.loading
        if ld.load_lbs is not None:
            _bold_para(doc, "Load", f"{ld.load_lbs} lbs")
        if ld.prior_load_lbs is not None:
            _bold_para(doc, "Prior", f"{ld.prior_load_lbs} lbs")
        if ld.reps_achieved:
            _bold_para(doc, "Reps achieved", ld.reps_achieved)
        if ld.progression_target:
            _bold_para(doc, "Progression target", ld.progression_target)

    for label, items in [("Cues", ex.cues), ("Regressions", ex.regressions), ("Progressions", ex.progressions)]:
        if items:
            p = doc.add_paragraph()
            p.add_run(f"{label}:").bold = True
            for item in items:
                doc.add_paragraph(item, style="List Bullet")


def _add_block(doc: Document, block: Block) -> None:
    title = f"{block.title} ({block.block_type})"
    if block.time_minutes:
        title += f" ~{block.time_minutes} min"
    doc.add_heading(title, level=2)

    if block.format:
        _bold_para(doc, "Format", block.format)

    for ex in block.exercises:
        _add_exercise(doc, ex)


def export(plan: TrainingSessionPlan) -> Path:
    doc = Document()
    meta = plan.meta

    doc.add_heading("Training Session Plan", level=1)

    for label, value in [
        ("Client", meta.client_name or "—"),
        ("Date", meta.session_date or "—"),
        ("Session #", str(meta.session_number or "—")),
        ("Duration", f"{meta.duration_minutes} min"),
        ("Focus", meta.focus),
        ("Constraints", ", ".join(meta.constraints) if meta.constraints else "—"),
    ]:
        _bold_para(doc, label, value)

    if plan.equipment_used:
        doc.add_heading("Equipment Used", level=2)
        for item in plan.equipment_used:
            doc.add_paragraph(item, style="List Bullet")

    for block in plan.blocks:
        _add_block(doc, block)

    if plan.progression_notes:
        doc.add_heading("Progression Notes (Next Session)", level=2)
        for n in plan.progression_notes:
            doc.add_paragraph(n, style="List Bullet")

    if plan.coaching_notes:
        doc.add_heading("Global Coaching Notes", level=2)
        for n in plan.coaching_notes:
            doc.add_paragraph(n, style="List Bullet")

    path = _output_path(plan)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
